from contextlib import asynccontextmanager
from fastapi import FastAPI, APIRouter, HTTPException, Request, Response, Cookie, Depends, Header
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import json
import uuid
import asyncio
import logging
import httpx
import jwt
import hashlib
import secrets
import re
import base64
import random
import io
import sys
import warnings
from docx import Document
from docx.shared import Pt
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from datetime import datetime, timezone, timedelta
from xml.sax.saxutils import escape as xml_escape
from starlette.responses import RedirectResponse

# --- Cache for Gemini responses ---
_gemini_cache = {}
_MAX_CACHE_ENTRIES = 100 # Limit cache size

def _generate_cache_key(model_name: str, system_instruction: str, prompt: str) -> str:
    """Generates a unique cache key for Gemini responses."""
    # Simple approach: combine model, system instruction, and prompt.
    # For more complex scenarios, consider hashing or more robust key generation.
    return f"{model_name}|{system_instruction}|{prompt}"

# --- Idempotency Cache ---
_IDEMPOTENCY_CACHE = {}
_IDEMPOTENCY_TTL = timedelta(seconds=600) # 10 minutes

def _clean_idempotency_cache():
    """Removes expired entries from the idempotency cache."""
    now = datetime.now(timezone.utc)
    keys_to_remove = [key for key, (data, expiry) in _IDEMPOTENCY_CACHE.items() if expiry < now]
    for key in keys_to_remove:
        del _IDEMPOTENCY_CACHE[key]

# --- Google SDK Initialization ---
class _GoogleSDKUnavailable(Exception):
    pass

genai = None
NotFound = _GoogleSDKUnavailable
InvalidArgument = _GoogleSDKUnavailable
GoogleBadRequest = _GoogleSDKUnavailable
VertexModel = None
GenerationConfig = None

if sys.version_info < (3, 14):
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", FutureWarning)
            import google.generativeai as genai
        from google.api_core.exceptions import BadRequest as GoogleBadRequest
        from google.api_core.exceptions import InvalidArgument, NotFound
        from vertexai.generative_models import GenerationConfig
        from vertexai.generative_models import GenerativeModel as VertexModel
    except Exception as e:
        logging.getLogger(__name__).warning(f"Google AI SDK imports unavailable: {e}")

# ============================================================
# INITIALIZATION & CONFIG
# ============================================================
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

ADMIN_EMAILS = set(e.strip().lower() for e in os.environ.get('ADMIN_EMAILS', 'bentaylors@hotmail.co.uk').split(',') if e.strip())

PAYPAL_CLIENT_ID     = os.environ.get('PAYPAL_CLIENT_ID', '')
PAYPAL_CLIENT_SECRET = os.environ.get('PAYPAL_CLIENT_SECRET', '')
PAYPAL_BASE_URL      = os.environ.get('PAYPAL_BASE_URL', 'https://api-m.sandbox.paypal.com')

_cors_env = os.environ.get('CORS_ORIGINS', '')
CORS_ORIGINS = [o.strip() for o in _cors_env.split(',') if o.strip()] or [
    "https://smartgiaoan.site",
    "https://www.smartgiaoan.site",
    "http://localhost:3000",
]
FRONTEND_URL = os.environ.get('FRONTEND_URL', CORS_ORIGINS[0]).rstrip("/")
BACKEND_PUBLIC_URL = os.environ.get('BACKEND_PUBLIC_URL', '').rstrip("/")

# ============================================================
# AI MODEL CONFIGURATION (OpenRouter Integration)
# ============================================================
# Free tier uses a free OpenRouter model.
# Basic tier uses OpenRouter's auto-selection.
# Premium tier uses OpenRouter with Claude 3 Opus for best performance.
OPENROUTER_MODEL_FREE    = "openrouter/free" # Or a specific free model like google/gemini-pro:free
OPENROUTER_MODEL_BASIC   = "openrouter/auto"
OPENROUTER_MODEL_PREMIUM = "openrouter/anthropic/claude-3-opus" # Best available model for premium users
GEMINI_MODEL_FREE = OPENROUTER_MODEL_FREE
GEMINI_MODEL_BASIC = OPENROUTER_MODEL_BASIC
GEMINI_MODEL_PREMIUM = OPENROUTER_MODEL_PREMIUM

TIER_CONFIG = {
    "free": {
        "model": GEMINI_MODEL_FREE,
        "lifetime_quota": 3,
        "monthly_quota": 3,          # FIX 4: added so billing/tier endpoint doesn't KeyError
        "ai_edits_per_month": 1,
        "has_word_editor": True,
        "has_ai_editor": False,
        "has_ads": True,
        "ad_frequency_base": 0.3,
    },
    "premium": {
        "model": GEMINI_MODEL_PREMIUM,
        "monthly_quota": 999999,
        "ai_edits_per_month": 999999,
        "has_word_editor": True,
        "has_ai_editor": True,
        "has_ads": False,
    },
    "pro": {
        "model": GEMINI_MODEL_PREMIUM,
        "monthly_quota": 999999,
        "ai_edits_per_month": 999999,
        "has_word_editor": True,
        "has_ai_editor": True,
        "has_ads": False,
    }
}

# ============================================================
# 🗺️ SMARTGIAOAN LOCALIZATION VAULT (v2.8)
# ============================================================
VIETNAM_CITIES = [
    "Hanoi", "Hai Phong", "Ha Long", "Sapa", "Ninh Binh",
    "Dien Bien", "Ha Giang", "Da Nang", "Hue", "Hoi An",
    "Nha Trang", "Da Lat", "Quy Nhon", "Dong Hoi",
    "Ho Chi Minh City", "Vung Tau", "Mui Ne", "Phu Quoc",
    "Can Tho", "Ben Tre"
]

VIETNAM_LANDMARKS = [
    "Ha Long Bay", "Fansipan", "Phong Nha Caves", "Mekong Delta",
    "Hoan Kiem Lake", "Temple of Literature", "Dragon Bridge",
    "Golden Bridge", "Ben Thanh Market", "Notre Dame Cathedral",
    "Landmark 81", "Cu Chi Tunnels", "the local street market",
    "the neighborhood pagoda", "the rice paddies"
]

VIETNAM_FOODS = [
    "Pho", "Banh Mi", "Bun Cha", "Spring Rolls", "Banh Xeo",
    "Egg Coffee", "Iced Milk Coffee", "dragonfruit",
    "rambutan", "mango"
]

ACTIVITY_VAULT = [
    "Code Breaker (substitution cipher or number-to-letter puzzle)",
    "Detective Story (fill-in-the-blank narrative with clues)",
    "Drawing Prompt (read and draw, or label-and-color)",
    "Categorization (sort words into buckets: food, place, animal, etc.)",
    "Matching (draw lines between columns: word-to-picture or word-to-definition)",
    "True or False (reading comprehension statements)",
    "Word Scramble (unscramble letters to form target vocabulary)",
    "Crossword (simple grid with picture clues for younger kids)",
    "Missing Letter (fill in the blank: b_n_n_ for 'banana')",
    "Odd One Out (circle the word that doesn't belong)",
]

# ============================================================
# MONGODB
# ============================================================
mongo_url = os.environ.get('MONGO_URL', '')
db_name   = os.environ.get('DB_NAME', 'smartgiaoan')
mongo_client = None
db = None

try:
    mongo_client = AsyncIOMotorClient(mongo_url, serverSelectionTimeoutMS=5000)
    db = mongo_client[db_name]
    logger.info("MongoDB client created successfully")
except Exception as e:
    logger.warning(f"MongoDB connection failed: {e}. Falling back to in-memory mock.")
    try:
        import mongomock
        mongo_client = mongomock.MongoClient()
        db = mongo_client[db_name]
        logger.info("Using mongomock in-memory MongoDB")
    except Exception as mock_err:
        logger.error(f"Failed to initialize mongomock: {mock_err}")
        raise

@asynccontextmanager
async def lifespan(app: FastAPI):
    yield
    if mongo_client:
        mongo_client.close()
        logger.info("MongoDB connection closed")

# ============================================================
# APP SETUP
# ============================================================
app = FastAPI(title="SmartGiaoAn API", version="3.3.0", lifespan=lifespan)
api_router = APIRouter(prefix="/api")

def hash_password(password: str) -> str:
    salt = secrets.token_hex(16)
    h = hashlib.pbkdf2_hmac('sha256', password.encode(), salt.encode(), 100000).hex()
    return f"{salt}:{h}"

def verify_password(password: str, hashed: str) -> bool:
    try:
        salt, stored = hashed.split(':', 1)
        h = hashlib.pbkdf2_hmac('sha256', password.encode(), salt.encode(), 100000).hex()
        return secrets.compare_digest(h, stored)
    except Exception:
        return False

# ============================================================
# PYDANTIC MODELS
# ============================================================
class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    user_id: str
    email: str
    name: str
    role: str = "Teacher"
    picture: Optional[str] = ""
    teaching_level: Optional[str] = ""
    class_size: Optional[str] = ""
    focus_area: Optional[str] = ""
    subscription_tier: str = "free"
    is_premium: bool = False
    free_used: int = 0
    bonus_credits: int = 0
    ai_edit_credits: int = 0
    monthly_reset_at: Optional[datetime] = None
    paypal_subscription_id: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class WorksheetRequest(BaseModel):
    level: str
    cefr: str
    skill: str
    topic: str
    num_questions: int = 24
    grammar_focus: Optional[str] = None


class SessionExchangeRequest(BaseModel):
    session_id: str

class EmailAuthRequest(BaseModel):
    email: str
    password: str
    name: Optional[str] = ""
    role: Optional[str] = "Teacher"
    heard_from: Optional[str] = ""

class RewardedAdRequest(BaseModel):
    tier: int
    reward_type: str = "worksheet"

class AIEditRequest(BaseModel):
    worksheet_id: str
    command: str

class UpdateWorksheetRequest(BaseModel):
    title: Optional[str] = None
    content: Optional[dict] = None
    is_public: Optional[bool] = None

class ProfileUpdateRequest(BaseModel):
    teaching_level: Optional[str] = None
    class_size: Optional[str] = None
    focus_area: Optional[str] = None
    name: Optional[str] = None
    role: Optional[str] = None

class LibraryUploadRequest(BaseModel):
    title: str
    description: str
    level: str
    skills: list[str] = Field(default_factory=list)
    topic: Optional[str] = ""
    is_public: bool = True

class WorksheetFixRequest(BaseModel):
    worksheetId: str
    originalPrompt: Optional[str] = ""
    feedback: str

class PayPalCaptureRequest(BaseModel):
    order_id: str
    product_type: str

# ============================================================
# AUTH HELPERS
# ============================================================
def _now():
    return datetime.now(timezone.utc)

def _parse_dt(v):
    if isinstance(v, datetime):
        return v
    dt = datetime.fromisoformat(str(v).replace('Z', '+00:00'))
    return dt if dt.tzinfo else dt.replace(tzinfo=timezone.utc)

async def _load_user(doc: dict) -> Optional[User]:
    if not doc:
        return None
    if "subscription_tier" not in doc:
        doc["subscription_tier"] = "premium" if doc.get("is_premium") else "free"
    if "ai_edit_credits" not in doc:
        doc["ai_edit_credits"] = 0
    return User(**doc)

async def refresh_user_credits(user: User) -> User:
    now = _now()
    needs_reset = False
    if not user.monthly_reset_at:
        needs_reset = True
    else:
        try:
            needs_reset = _parse_dt(user.monthly_reset_at) < now
        except Exception:
            needs_reset = True

    if needs_reset:
        config = TIER_CONFIG.get(user.subscription_tier, TIER_CONFIG["free"])
        updates = {
            "ai_edit_credits": config["ai_edits_per_month"],
            "monthly_reset_at": (now + timedelta(days=30)).isoformat()
        }
        if user.subscription_tier != "free":
            updates["free_used"] = 0
            user.free_used = 0
        await db.users.update_one(
            {"user_id": user.user_id},
            {"$set": updates}
        )
        user.ai_edit_credits = config["ai_edits_per_month"]
        user.monthly_reset_at = now + timedelta(days=30)

    return user

async def _create_session(user_id: str, response: Response) -> str:
    token = str(uuid.uuid4())
    expires = _now() + timedelta(days=7)
    await db.user_sessions.insert_one({
        "user_id": user_id, "session_token": token,
        "expires_at": expires.isoformat(), "created_at": _now().isoformat(),
    })
    response.set_cookie(
        key="session_token", value=token,
        httponly=True, secure=True, samesite="none",
        max_age=7 * 24 * 3600, path="/"
    )
    return token

# ============================================================
# GEMINI ENGINE
# ============================================================
GEMINI_REGION   = os.environ.get('GEMINI_REGION', 'us-central1')
USE_VERTEX_AI   = False
_vertex_project = None

adc_json_raw = os.environ.get('GOOGLE_APPLICATION_CREDENTIALS_JSON', '').strip()
api_key      = os.environ.get('GEMINI_API_KEY', '').strip()
OPENROUTER_API_KEY = os.environ.get('OPENROUTER_API_KEY', '').strip()


def _parse_google_credentials_json(raw: str) -> dict:
    """Parse service-account JSON from Render env vars without logging secrets."""
    value = raw.strip()
    if not value:
        raise ValueError("GOOGLE_APPLICATION_CREDENTIALS_JSON is empty")

    if (value.startswith("'") and value.endswith("'")) or (value.startswith('"') and value.endswith('"')):
        value = value[1:-1].strip()

    candidates = [value]

    try:
        decoded = base64.b64decode(value, validate=True).decode("utf-8").strip()
        if decoded:
            candidates.append(decoded)
    except Exception:
        pass

    for candidate in candidates:
        try:
            parsed = json.loads(candidate)
            if isinstance(parsed, str):
                parsed = json.loads(parsed)
            if not isinstance(parsed, dict):
                raise ValueError("credential JSON did not decode to an object")
            return parsed
        except json.JSONDecodeError:
            fixed = re.sub(
                r'("private_key"\s*:\s*")(.*?)("\s*,\s*"client_email")',
                lambda m: m.group(1) + m.group(2).replace("\n", "\\n") + m.group(3),
                candidate,
                flags=re.DOTALL,
            )
            if fixed != candidate:
                try:
                    parsed = json.loads(fixed)
                    if isinstance(parsed, dict):
                        return parsed
                except json.JSONDecodeError:
                    pass

    raise ValueError(
        "GOOGLE_APPLICATION_CREDENTIALS_JSON is not valid service-account JSON. "
        "Paste the raw JSON object exactly, or set it to base64-encoded JSON."
    )

if adc_json_raw:
    try:
        creds_dict = _parse_google_credentials_json(adc_json_raw)
        adc_path = '/tmp/gcp_adc.json'
        with open(adc_path, 'w') as f:
            json.dump(creds_dict, f)
        os.chmod(adc_path, 0o600)
        os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = adc_path

        import google.auth
        _gcp_creds, _vertex_project = google.auth.default()

        if VertexModel and GenerationConfig:
            import vertexai
            vertexai.init(project=_vertex_project, location=GEMINI_REGION)
            USE_VERTEX_AI = True
            logger.info(f"AI Engine: Vertex AI | project={_vertex_project} | region={GEMINI_REGION} | geo-block bypassed ✓")
        else:
            logger.warning("Vertex AI credentials were provided, but Vertex AI SDK imports are unavailable.")
            if api_key and genai:
                genai.configure(api_key=api_key)
                logger.warning("AI Engine: API key fallback enabled.")

    except Exception as e:
        logger.error(f"Vertex AI init failed: {e}")
        if api_key and genai:
            genai.configure(api_key=api_key)
            logger.warning("AI Engine: API key fallback — change Render region to Oregon to avoid 400 location errors.")
        else:
            logger.error("CRITICAL: No usable Google credentials.")

elif api_key and genai:
    genai.configure(api_key=api_key)
    logger.warning(
        "AI Engine: API key mode. "
        "If you see '400 User location not supported', go to "
        "Render Dashboard → Settings → Region → Oregon (US West)."
    )
else:
    logger.error("CRITICAL: NO USABLE GOOGLE CREDENTIALS — all AI calls will use the disabled mock.")


def _is_location_error(exc: Exception) -> bool:
    msg = str(exc).lower()
    return "user location is not supported" in msg or ("failed_precondition" in msg and "location" in msg)


def _build_model(model_name: str, system_instruction: str):
    if USE_VERTEX_AI:
        return VertexModel(
            model_name,
            system_instruction=system_instruction,
            generation_config=GenerationConfig(
                response_mime_type="application/json",
                temperature=0.8,
            ),
        )
    elif genai: # Check if genai is not None
        return genai.GenerativeModel(
            model_name=model_name,
            system_instruction=system_instruction,
            generation_config={"response_mime_type": "application/json", "temperature": 0.8},
        )
    else:
        # If genai is None, return a mock or raise a specific error
        logger.error("AI features are disabled: No Google credentials found.")
        # Define mock classes locally to avoid polluting global scope if not needed
        class MockGenerateContentResponse:
            def __init__(self):
                self.text = "AI features are currently disabled due to missing credentials."
                self.candidates = []
            def __str__(self):
                return self.text

        class MockGenerativeModel:
            def __init__(self, model_name, system_instruction, generation_config):
                self.model_name = model_name
                self.system_instruction = system_instruction
                self.generation_config = generation_config
                logger.warning(f"MockGenerativeModel initialized for {model_name}")

            def start_chat(self, **kwargs):
                logger.warning("MockGenerativeModel.start_chat called - AI features disabled.")
                class MockChatSession:
                    def send_message(self, message, **kwargs):
                        logger.warning("MockChatSession.send_message called - AI features disabled.")
                        return MockGenerateContentResponse()
                return MockChatSession()

            def generate_content(self, prompt, **kwargs):
                logger.warning("MockGenerativeModel.generate_content called - AI features disabled.")
                return MockGenerateContentResponse()
        
        return MockGenerativeModel(model_name, system_instruction, {"response_mime_type": "application/json", "temperature": 0.8})


def build_system_prompt(level: str, skill: str = "", topic: str = "", num_questions: int = 24) -> str:
    selected_cities     = random.sample(VIETNAM_CITIES,     k=2)
    selected_landmarks  = random.sample(VIETNAM_LANDMARKS,  k=2)
    selected_foods      = random.sample(VIETNAM_FOODS,      k=3)
    selected_activities = random.sample(ACTIVITY_VAULT,     k=4)

    activity_directive = "\n".join(
        f"    {i+1}. {act}" for i, act in enumerate(selected_activities)
    )

    if level in ("Kindergarten", "Primary"):
        formatting_rules = f"""
FORMAT RULES (Kindergarten / Primary — MANDATORY):
- You are generating a worksheet for YOUNG LEARNERS (ages 4-11).
- STRICTLY DISABLE the 3-Page Rule. Do NOT force Reading + Vocab + Grammar + Writing sections.
- The entire document MUST fit comfortably on 1 to 2 A4 pages. White space is your friend.
- Use LARGE, friendly fonts. Short sentences. Lots of visual breathing room.
- Every activity must be achievable with a pencil and crayons; no essay writing.
- If you include a reading passage, keep it under 40 words for Kindergarten, under 80 for Primary.
- NUMBER OF ITEMS: {num_questions} (honour this exactly — never pad, never exceed 32)
- Kindergarten: count tasks, not questions. 6-10 tasks = full worksheet.
"""
    else:
        formatting_rules = f"""
FORMAT RULES (Secondary / IELTS — MANDATORY):
- You MUST follow the strict 3-Page Rule:
    Page 1: Reading Comprehension (passage + questions)
    Page 2: Vocabulary & Grammar (exercises based on the reading)
    Page 3: Writing Task (structured output: essay, letter, or long-form response)
- The tone is academic and age-appropriate for teenagers or adult test-takers.
- NUMBER OF ITEMS: {num_questions} (honour this exactly — never pad, never exceed 32)
- IELTS: one full passage + question set (15-20 items) is one complete worksheet.
"""

    listening_rule = (
        "LISTENING WORKSHEET RULE: You MUST include a listening_script field at the top level "
        "of the JSON. This is the full text of the audio track the teacher will read aloud. "
        "Design all questions to test comprehension of that script specifically."
        if skill and skill.lower() == "listening" else ""
    )

    prompt = f"""You are SmartGiaoAn, an expert Vietnamese ESL worksheet designer.

LOCALIZED CONTEXT FOR THIS WORKSHEET (use these naturally in activities):
- Cities: {selected_cities[0]} and {selected_cities[1]}
- Landmarks: {selected_landmarks[0]} and {selected_landmarks[1]}
- Foods: {selected_foods[0]}, {selected_foods[1]}, and {selected_foods[2]}

{formatting_rules}

TOPIC: {topic}

DYNAMIC GAME DIRECTOR (MANDATORY — DO NOT USE A STATIC FORMULA):
You must select EXACTLY 3 to 4 completely different activity types for this worksheet.
For THIS specific generation, the Director has chosen:
{activity_directive}

RULES FOR THE DIRECTOR:
- You MUST use the activity types listed above. Do not swap them out.
- Each activity must feel distinct. Do not repeat the same mechanic twice.
- Weave the localized cities, landmarks, and foods into the activities naturally.
- For Kindergarten/Primary: favor drawing, matching, coloring, and simple puzzles.
- For Secondary/IELTS: you may adapt the activities to be more analytical (e.g., a Detective Story becomes a reading-comprehension mystery).

UNSPLASH HEADER IMAGE REQUIREMENT:
- At the top of your JSON output, include a field named "header_image_query".
- The value must be a specific, photorealistic Unsplash search query for a 2K landscape of Vietnamese nature or cityscape based on one of the injected locations.
- Examples: "Da Lat pine forest misty morning photorealistic", "Ha Long Bay limestone karsts sunset 4k", "Mekong Delta floating market aerial".
- STRICTLY NO CARTOONS, NO CLIP ART, NO ILLUSTRATIONS. Real photography only.

ANSWER KEY ISOLATION (CRITICAL FOR PRINT ENGINE):
- ALL answers, solutions, teacher notes, and correct mappings MUST be placed inside a top-level JSON key named "answer_key".
- The "answer_key" object must contain clear, labeled answers for every activity you generated.
- The frontend print CSS uses the ".answer-key-section" class to force answers onto a separate A4 page. Your JSON structure must support this by keeping answers out of the main "activities" array.
- Do NOT mix answers into the student-facing activity content.

OUTPUT FORMAT — VALID JSON ONLY:
Return a single, valid JSON object. No markdown code fences. No commentary outside the JSON.
The structure must be:

{{
  "header_image_query": "string — specific Unsplash query for 2K photorealistic Vietnamese landscape",
  "title": "string — catchy, localized title using one of the cities or landmarks",
  "level": "e.g. {level}",
  "audience": "e.g. Primary, age 9-12",
  "skill_focus": "{skill or 'mixed'}",
  "topic": "{topic}",
  "learning_objectives": ["objective 1", "objective 2", "objective 3"],
  "vocab_list": [
    {{"word": "...", "definition": "...", "example": "Vietnamese-localised example sentence"}}
  ],
  "listening_script": "ONLY present if skill is listening. Full read-aloud script here.",
  "sections": [
    {{
      "section_number": 1,
      "section_title": "Part 1: ...",
      "section_type": "passage | vocabulary | grammar | comprehension | listening | activity",
      "instruction": "Clear student-facing instruction.",
      "content": "Passage text, word list, or other non-question content. Omit key if not applicable.",
      "items": [
        {{
          "number": 1,
          "question": "The question or task stem",
          "options": ["A. ...", "B. ...", "C. ...", "D. ..."],
          "answer_line": "_______________"
        }}
      ]
    }}
  ],
  "writing_task": {{
    "prompt": "The full writing task instruction",
    "success_criteria": ["criterion 1", "criterion 2", "criterion 3"]
  }},
  "answer_key": [
    {{"number": 1, "answer": "correct answer"}}
  ],
  "teacher_notes": [
    "Note 1: a common Vietnamese L1 interference error relevant to this grammar or skill",
    "Note 2: a pronunciation or spelling trap for Vietnamese speakers",
    "Note 3: a cultural connection point or suggested board activity",
    "Note 4: differentiation tip for mixed-ability classes",
    "Note 5: suggested follow-up activity"
  ],
  "extension_activity": "One optional challenge task for fast finishers. Must use the same topic and push one level higher."
}}

{listening_rule}

CONTENT GUIDELINES:
- "content" should be flexible per activity type. For example:
    - matching: {{"left": ["...", "..."], "right": ["...", "..."]}}
    - true_false: {{"statements": ["...", "..."]}}
    - word_scramble: {{"words": ["...", "..."]}}
    - categorization: {{"categories": ["...", "..."], "items": ["...", "..."]}}
    - drawing_prompt: {{"prompt": "...", "vocabulary": ["...", "..."]}}
- Keep vocabulary aligned with the topic AND the localized context.
- Ensure the JSON is minified or pretty-printed, but ALWAYS syntactically valid.
- Vietnamese names (rotate freely): Minh, Lan, Huy, Trang, Nam, Linh, Duc, Mai, Khoa, Phuong, An, Bao, Chi, Dung, Ha, Khanh, Long, My, Nhi, Quang, Son, Thu, Tuan, Vy
- Vietnamese culture & food: Tet, Mid-Autumn Festival, ao dai, dong ho paintings, water puppetry, pho, banh mi, bun bo Hue, com tam, banh xeo, che, ca phe trung, xe om, motorbike culture, family-centred values, ancestor worship, five-fruit tray
"""
    return prompt


async def _run_gemini(prompt: str, level: str, model_name: str, skill: str = "", topic: str = "", num_questions: int = 24) -> dict:
    system_instruction = build_system_prompt(level, skill=skill, topic=topic, num_questions=num_questions)
    cache_key = _generate_cache_key(model_name, system_instruction, prompt)

    # --- Cache Lookup ---
    if cache_key in _gemini_cache:
        logger.info(f"Cache hit for prompt: {prompt[:50]}...")
        return _gemini_cache[cache_key]

    seen: set = set()
    model_chain = [
        m for m in ([model_name] + _GEMINI_FALLBACKS.get(model_name, ["gemini-1.5-flash"]))
        if not (m in seen or seen.add(m))
    ]

    last_error = "Unknown error"

    for current_model in model_chain:
        logger.info(
            f"[Gemini] model={current_model} | "
            f"engine={'vertex/' + GEMINI_REGION if USE_VERTEX_AI else 'genai/api-key'}"
        )

        for attempt in range(3):
            try:
                if current_model.startswith("openrouter/") or ":free" in current_model:
                    # BRANCH A: OpenRouter via HTTPX
                    logger.info(f"[OpenRouter] Routing request to {current_model}")
                    headers = {
                        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                        "HTTP-Referer": "https://smartgiaoan.site",
                        "Content-Type": "application/json"
                    }
                    payload = {
                        "model": current_model,
                        "messages": [
                            {"role": "system", "content": system_instruction},
                            {"role": "user", "content": prompt}
                        ],
                        "response_format": {"type": "json_object"}
                    }

                    async with httpx.AsyncClient(timeout=90.0) as client:
                        resp = await client.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload)
                        resp.raise_for_status()
                        data = resp.json()
                        raw = data['choices'][0]['message']['content'].strip()
                else:
                    # BRANCH B: Existing Google/Vertex AI SDK
                    model = _build_model(current_model, system_instruction)
                    result = await asyncio.wait_for(
                        asyncio.to_thread(model.generate_content, prompt),
                        timeout=90.0
                    )
                    if not result.candidates:
                        feedback = getattr(result, 'prompt_feedback', None)
                        block_reason = getattr(feedback, 'block_reason', 'Unknown') if feedback else 'Unknown'
                        last_error = f"Content blocked: {block_reason}"
                        logger.warning(f"[Gemini] {last_error}")
                        await asyncio.sleep(1)
                        continue
                    raw = result.text.strip()

                # --- SHARED JSON PARSING (Keep your existing regex/json parsing here) ---
                if not raw:
                    last_error = "Empty response"
                    logger.warning(f"[AI] {last_error} (attempt {attempt+1})")
                    await asyncio.sleep(1)
                    continue

                raw = re.sub(r'^(?:json)?\s*', '', raw, flags=re.IGNORECASE)
                raw = re.sub(r'\s*$', '', raw).strip()
                parsed = json.loads(raw)

                # --- Cache Storage ---
                if len(_gemini_cache) >= _MAX_CACHE_ENTRIES:
                    # Simple eviction: remove the oldest entry (first one added)
                    oldest_key = next(iter(_gemini_cache))
                    del _gemini_cache[oldest_key]
                    logger.debug(f"Cache full, evicted oldest entry: {oldest_key[:50]}...")
                
                _gemini_cache[cache_key] = parsed
                logger.info(f"Cached response for prompt: {prompt[:50]}...")

                if current_model != model_name:
                    logger.warning(f"[AI] Used fallback '{current_model}' (primary '{model_name}' was unavailable)")

                return parsed

            except httpx.HTTPStatusError as e:
                last_error = f"OpenRouter HTTP error: {e}"
                logger.error(f"[OpenRouter] {last_error} (attempt {attempt+1})")
                if e.response.status_code == 429:
                    logger.warning("[OpenRouter] Rate limited; trying next model/provider.")
                    break
                await asyncio.sleep(1)

            except httpx.RequestError as e:
                last_error = f"OpenRouter request error: {e}"
                logger.error(f"[OpenRouter] {last_error} (attempt {attempt+1})")
                await asyncio.sleep(1)

            except GoogleBadRequest as e:
                if _is_location_error(e):
                    raise HTTPException(status_code=503, detail=(
                        "AI Engine: Google blocked this server's location (400 FAILED_PRECONDITION). "
                        "IMMEDIATE FIX: Render Dashboard → Settings → Region → 'Oregon (US West)'. "
                        "PERMANENT FIX: Add GOOGLE_APPLICATION_CREDENTIALS_JSON env var to use Vertex AI routing."
                    ))
                last_error = f"Bad request: {e}"
                logger.error(f"[Gemini] {last_error} (attempt {attempt+1})")

            except Exception as e:
                if _is_location_error(e):
                    raise HTTPException(status_code=503, detail=(
                        "AI Engine: Google blocked this server's location (400 FAILED_PRECONDITION). "
                        "IMMEDIATE FIX: Render Dashboard → Settings → Region → 'Oregon (US West)'. "
                        "PERMANENT FIX: Add GOOGLE_APPLICATION_CREDENTIALS_JSON env var to use Vertex AI routing."
                    ))
                if isinstance(e, (NotFound, InvalidArgument)):
                    last_error = f"Model not found: {e}"
                    logger.info(f"[Gemini] {last_error} — trying fallback")
                    break
                if isinstance(e, asyncio.TimeoutError):
                    last_error = "Timed out after 90s"
                    logger.error(f"[Gemini] {last_error} (attempt {attempt+1})")
                elif isinstance(e, json.JSONDecodeError):
                    last_error = f"JSON parse error: {e}"
                    logger.error(f"[Gemini] {last_error} (attempt {attempt+1})")
                else:
                    last_error = f"API error: {e}"
                    logger.error(f"[Gemini] {last_error} (attempt {attempt+1})")

            await asyncio.sleep(1)

    raise HTTPException(
        status_code=500,
        detail=f"AI Engine failed: all models exhausted. Last error: {last_error}"
    )

# ============================================================
# API ROUTES — AUTH
# ============================================================
# JWT secret for email verification
EMAIL_VERIFICATION_JWT_SECRET = os.environ.get('EMAIL_VERIFICATION_JWT_SECRET') or os.environ.get('JWT_VERIFICATION_SECRET')
if not EMAIL_VERIFICATION_JWT_SECRET:
    logger.warning("EMAIL_VERIFICATION_JWT_SECRET not configured. Email verification will not work.")

# Add email_verified field to User model if it doesn't exist
# This is handled by Pydantic's ConfigDict(extra="ignore") and default values in _load_user

# Add email_verified to the User model definition if it's not already there
# (Assuming it's already added or handled by default in _load_user)

# Add the /auth/send-verification endpoint
@api_router.post("/auth/send-verification")
async def send_verification(user: User = Depends(require_user)):
    if getattr(user, 'email_verified', False):
        return {"status": "already_verified"}
    
    if not EMAIL_VERIFICATION_JWT_SECRET:
        raise HTTPException(status_code=500, detail="Email verification is not configured on the server.")

    try:
        exp = datetime.utcnow() + timedelta(hours=24) # Token valid for 24 hours
        data = {'user_id': user.user_id, 'email': user.email, 'exp': exp.isoformat()}
        token = jwt.encode(data, EMAIL_VERIFICATION_JWT_SECRET, algorithm='HS256')
        
        # Construct verification link
        link = f"{FRONTEND_URL}/verify-email?token={token}"
        
        await _send_email(user.email, 'Verify your SmartGiaoAn email', f'Please verify your email by clicking this link: {link}')
        
        return {"status": "verification_sent"}
    except Exception as e:
        logger.error(f"Error sending verification email: {e}")
        raise HTTPException(status_code=500, detail="Failed to send verification email.")

# Add the /auth/verify-email endpoint
class EmailVerificationRequest(BaseModel):
    token: str

@api_router.post("/auth/verify-email")
async def verify_email(payload: EmailVerificationRequest):
    if not EMAIL_VERIFICATION_JWT_SECRET:
        raise HTTPException(status_code=500, detail="Email verification is not configured on the server.")
    
    try:
        data = jwt.decode(payload.token, EMAIL_VERIFICATION_JWT_SECRET, algorithms=["HS256"])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=400, detail="Verification token has expired.")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=400, detail="Invalid verification token.")
    
    user_id = data.get("user_id")
    if not user_id:
        raise HTTPException(status_code=400, detail="Invalid token payload: missing user_id.")
    
    # Update user's email_verified status
    result = await db.users.update_one({"user_id": user_id}, {"$set": {"email_verified": True}})
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="User not found.")
    if result.modified_count == 0:
        # This could happen if the user was already verified, or if the update failed for some reason.
        # We can consider it a success if the user exists and is already verified.
        user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
        if user_doc and user_doc.get("email_verified"):
            return {"status": "verified"}
        else:
            raise HTTPException(status_code=500, detail="Failed to update verification status.")

    return {"status": "verified"}
async def get_current_user_optional(
    request: Request,
    session_token: Optional[str] = Cookie(None),
    authorization: Optional[str] = Header(None)
) -> Optional[User]:
    token = session_token
    if authorization and authorization.lower().startswith("bearer "):
        token = authorization.split(" ", 1)[1].strip()
    if not token:
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.lower().startswith("bearer "):
            token = auth_header.split(" ", 1)[1].strip()
    if not token:
        return None
    session = await db.user_sessions.find_one({"session_token": token}, {"_id": 0})
    if not session or _parse_dt(session["expires_at"]) < _now():
        return None
    doc = await db.users.find_one({"user_id": session["user_id"]}, {"_id": 0})
    return await _load_user(doc)


async def require_user(user: Optional[User] = Depends(get_current_user_optional)) -> User:
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated or session expired")
    return user


@api_router.post("/auth/register")
async def auth_register(payload: EmailAuthRequest, response: Response):
    email = payload.email.strip().lower()
    if await db.users.find_one({"email": email}):
        raise HTTPException(status_code=400, detail="Email already registered")
    user_id  = f"user_{uuid.uuid4().hex[:12]}"
    is_admin = email in ADMIN_EMAILS
    doc = {
        "user_id": user_id, "email": email,
        "name": payload.name or email.split("@")[0],
        "role": payload.role or "Teacher",
        "password_hash": hash_password(payload.password),
        "is_premium": is_admin,
        "subscription_tier": "premium" if is_admin else "free",
        "free_used": 0, "bonus_credits": 0, "ai_edit_credits": 0,
        "created_at": _now().isoformat()
    }
    await db.users.insert_one(doc)
    token = await _create_session(user_id, response)
    return {"user": {k: v for k, v in doc.items() if k not in ["_id", "password_hash"]}, "session_token": token}

@api_router.post("/auth/login")
async def auth_login(payload: EmailAuthRequest, response: Response):
    email = payload.email.strip().lower()
    doc   = await db.users.find_one({"email": email})
    if not doc or not doc.get("password_hash") or not verify_password(payload.password, doc["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    token = await _create_session(doc["user_id"], response)
    return {"user": {k: v for k, v in doc.items() if k not in ["_id", "password_hash"]}, "session_token": token}

@api_router.post("/auth/session")
async def auth_session(payload: SessionExchangeRequest):
    session = await db.user_sessions.find_one({"session_token": payload.session_id}, {"_id": 0})
    if not session or _parse_dt(session["expires_at"]) < _now():
        raise HTTPException(status_code=401, detail="Invalid or expired session")
    doc = await db.users.find_one({"user_id": session["user_id"]}, {"_id": 0})
    user = await _load_user(doc)
    if not user:
        raise HTTPException(status_code=401, detail="Session user not found")
    user = await refresh_user_credits(user)
    return {"user": user.model_dump(), "session_token": payload.session_id}

# New endpoint for Google OAuth callback
@api_router.get("/auth/google-callback")
async def google_oauth_callback(
    request: Request,
    response: Response,
    code: Optional[str] = None,
    state: Optional[str] = None # Note: Frontend does not currently send a 'state' parameter for CSRF protection.
):
    if not code:
        raise HTTPException(status_code=400, detail="Missing authorization code")

    # Retrieve Google API credentials from environment variables
    google_client_id = os.environ.get('GOOGLE_CLIENT_ID')
    google_client_secret = os.environ.get('GOOGLE_CLIENT_SECRET')
    if not google_client_id or not google_client_secret:
        logger.error("Google OAuth credentials not configured. Missing GOOGLE_CLIENT_ID or GOOGLE_CLIENT_SECRET.")
        raise HTTPException(status_code=500, detail="Server configuration error: Google OAuth credentials missing.")

    dynamic_redirect_uri = (
        f"{BACKEND_PUBLIC_URL}/api/auth/google-callback"
        if BACKEND_PUBLIC_URL
        else str(request.url_for("google_oauth_callback"))
    )

    token_url = "https://oauth2.googleapis.com/token"
    userinfo_url = "https://www.googleapis.com/oauth2/v3/userinfo"

    try:
        # 1. Exchange authorization code for tokens
        token_response = await httpx.AsyncClient().post(
            token_url,
            data={
                "grant_type": "authorization_code",
                "code": code,
                "client_id": google_client_id,
                "client_secret": google_client_secret,
                "redirect_uri": dynamic_redirect_uri,
            }
        )
        token_response.raise_for_status()
        tokens = token_response.json()
        access_token = tokens.get("access_token")
        id_token = tokens.get("id_token") # id_token contains user info

        if not access_token or not id_token:
            logger.error(f"Failed to get access_token or id_token from Google: {tokens}")
            raise HTTPException(status_code=500, detail="Failed to obtain tokens from Google.")

        # 2. Get user info using access token
        userinfo_response = await httpx.AsyncClient().get(
            userinfo_url,
            headers={"Authorization": f"Bearer {access_token}"}
        )
        userinfo_response.raise_for_status()
        user_info = userinfo_response.json()

        email = user_info.get("email")
        name = user_info.get("name")
        picture = user_info.get("picture")

        if not email:
            logger.error(f"User info from Google missing email: {user_info}")
            raise HTTPException(status_code=500, detail="Could not retrieve user email from Google.")

        # 3. Find or create user in the database
        user_doc = await db.users.find_one({"email": email})
        is_admin = email in ADMIN_EMAILS

        if not user_doc:
            user_id = f"user_{uuid.uuid4().hex[:12]}"
            user_doc = {
                "user_id": user_id,
                "email": email,
                "name": name or email.split("@")[0],
                "picture": picture or "",
                "role": "Teacher", # Default role
                "is_premium": is_admin,
                "subscription_tier": "premium" if is_admin else "free",
                "free_used": 0,
                "bonus_credits": 0,
                "ai_edit_credits": 0,
                "created_at": _now().isoformat()
            }
            await db.users.insert_one(user_doc)
            logger.info(f"New user created: {email}")
        else:
            # Update user info if changed (e.g., picture, name)
            updates = {}
            if user_doc.get("name") != name and name:
                updates["name"] = name
            if user_doc.get("picture") != picture and picture:
                updates["picture"] = picture
            if updates:
                await db.users.update_one({"user_id": user_doc["user_id"]}, {"$set": updates})

        # 4. Create session and set cookie
        token = await _create_session(user_doc["user_id"], response)

        frontend_origin = state.rstrip("/") if state else FRONTEND_URL
        if frontend_origin not in CORS_ORIGINS:
            frontend_origin = FRONTEND_URL
        frontend_redirect_url = f"{frontend_origin}/auth/callback?session_id={token}"
        return RedirectResponse(url=frontend_redirect_url, status_code=303)

    except HTTPException:
        raise # Re-raise HTTPException to be handled by FastAPI
    except httpx.HTTPStatusError as e:
        logger.error(f"HTTP error during Google OAuth: {e.response.status_code} - {e.response.text}")
        raise HTTPException(status_code=500, detail=f"Google OAuth error: {e.response.status_code}")
    except Exception as e:
        logger.error(f"An unexpected error occurred during Google OAuth: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {e}")

@api_router.get("/auth/me")
async def auth_me(user: User = Depends(require_user)):
    return (await refresh_user_credits(user)).model_dump()

@api_router.post("/auth/logout")
async def auth_logout(response: Response, request: Request):
    token = request.cookies.get("session_token")
    auth  = request.headers.get("Authorization", "")
    if auth.lower().startswith("bearer "):
        token = auth.split(" ")[1]
    if token:
        await db.user_sessions.delete_many({"session_token": token})
    response.delete_cookie("session_token", path="/")
    return {"status": "logged_out"}

@api_router.get("/auth/export")
async def export_account(user: User = Depends(require_user)):
    worksheets = await db.worksheets.find({"user_id": user.user_id}, {"_id": 0}).to_list(None)
    return {"user": user.model_dump(), "worksheets": worksheets}

@api_router.delete("/auth/delete-account")
async def delete_account(user: User = Depends(require_user)):
    await db.users.delete_one({"user_id": user.user_id})
    await db.worksheets.delete_many({"user_id": user.user_id})
    await db.user_sessions.delete_many({"user_id": user.user_id})
    return {"status": "deleted"}

@api_router.put("/auth/profile")
async def update_profile(payload: ProfileUpdateRequest, user: User = Depends(require_user)):
    allowed = {
        "teaching_level": payload.teaching_level,
        "class_size": payload.class_size,
        "focus_area": payload.focus_area,
        "name": payload.name,
        "role": payload.role,
    }
    updates = {k: v.strip() if isinstance(v, str) else v for k, v in allowed.items() if v is not None}
    if updates:
        updates["updated_at"] = _now().isoformat()
        await db.users.update_one({"user_id": user.user_id}, {"$set": updates})
    doc = await db.users.find_one({"user_id": user.user_id}, {"_id": 0})
    refreshed = await _load_user(doc)
    return refreshed.model_dump()

# ============================================================
# API ROUTES — WORKSHEETS
# ============================================================
@api_router.post("/worksheets/generate")
async def generate_ws(payload: WorksheetRequest, user: User = Depends(require_user), request: Request = Depends()): # Added Request dependency
    # --- Idempotency Check ---
    idempotency_key = request.headers.get("Idempotency-Key")
    if idempotency_key:
        _clean_idempotency_cache() # Clean up expired entries
        if idempotency_key in _IDEMPOTENCY_CACHE:
            cached_data, expiry = _IDEMPOTENCY_CACHE[idempotency_key]
            if expiry >= _now():
                logger.info(f"Idempotency cache hit for key: {idempotency_key}")
                return cached_data # Return cached result
            else:
                logger.info(f"Idempotency cache expired for key: {idempotency_key}")
                del _IDEMPOTENCY_CACHE[idempotency_key] # Remove expired entry

    user   = await refresh_user_credits(user)
    config = TIER_CONFIG.get(user.subscription_tier, TIER_CONFIG["free"])

    should_show_sponsor = False
    sponsor_duration    = 0
    if user.subscription_tier == "free" and user.free_used > 0:
        prob = config.get("ad_frequency_base", 0.3) + min(user.free_used / 10, 1.0) * 0.4
        should_show_sponsor = random.random() < prob
        if should_show_sponsor:
            sponsor_duration = random.choice([15, 30, 60])

    prompt = (
        f"Create one complete, print-ready ESL worksheet using these exact specifications.\n\n"
        f"LEVEL: {payload.level} (CEFR {payload.cefr})\n"
        f"SKILL: {payload.skill}\n"
        f"TOPIC: {payload.topic}\n"
        f"GRAMMAR FOCUS: {payload.grammar_focus or 'Choose the most appropriate grammar point for this level and topic'}\n"
        f"NUMBER OF ITEMS: {payload.num_questions} (honour this exactly — never pad, never exceed 32)\n\n"
        "STRICT RULES:\n"
        f"1. LOCALISATION — Localise everything to Vietnam by default. Use Vietnamese names, places, food, and culture naturally throughout the passage and all example sentences.\n"
        f"2. LEVEL CEILING — Every word, sentence, and instruction must be strictly within the {payload.cefr} CEFR vocabulary ceiling.\n"
        "3. READING PASSAGE — If applicable, the passage must read like a real story or authentic text with a character, a setting, and an event. Never a list of facts.\n"
        f"4. GRAMMAR IN CONTEXT — Weave '{payload.grammar_focus or 'the chosen grammar point'}' into the passage and practice sections naturally.\n"
        "5. ANSWER KEY — The answer_key array must contain the correct answer for every single numbered item without exception.\n"
        "6. TEACHER NOTES — Must reference at least one Vietnamese L1 interference error specific to this grammar point or skill.\n"
        "7. OUTPUT — Return ONLY the raw JSON matching the schema in your instructions. No markdown. No code fences. No preamble."
    )

    ws_data = await _run_gemini(prompt, payload.level, model_name=config["model"], skill=payload.skill, topic=payload.topic, num_questions=payload.num_questions)
    ws_id   = f"ws_{uuid.uuid4().hex[:12]}"
    ws_doc  = {
        "worksheet_id": ws_id, "user_id": user.user_id,
        "title": f"{payload.topic} - {payload.skill}",
        "level": payload.level, "cefr": payload.cefr,
        "skill": payload.skill, "topic": payload.topic,
        "content": ws_data, "is_public": True, "created_at": _now().isoformat()
    }
    await db.worksheets.insert_one(ws_doc)
    await db.users.update_one({"user_id": user.user_id}, {"$inc": {"free_used": 1}})

    result = {k: v for k, v in ws_doc.items() if k != "_id"}

    # --- Store in Idempotency Cache ---
    if idempotency_key:
        expiry_time = _now() + _IDEMPOTENCY_TTL
        _IDEMPOTENCY_CACHE[idempotency_key] = (result, expiry_time)
        logger.info(f"Stored result in idempotency cache for key: {idempotency_key}")

    if should_show_sponsor:
        result["show_sponsor"] = True
        result["sponsor_duration"] = sponsor_duration
    return result

@api_router.get("/worksheets")
async def list_ws(user: User = Depends(require_user)):
    return await db.worksheets.find({"user_id": user.user_id}, {"_id": 0}).sort("created_at", -1).to_list(100)

@api_router.get("/worksheets/{worksheet_id}")
async def get_worksheet(worksheet_id: str, user: User = Depends(require_user)):
    ws = await db.worksheets.find_one({"worksheet_id": worksheet_id}, {"_id": 0})
    if not ws:
        raise HTTPException(status_code=404, detail="Worksheet not found")
    if ws.get("user_id") != user.user_id and not ws.get("is_public", False):
        raise HTTPException(status_code=403, detail="This worksheet is private")
    return ws

@api_router.get("/worksheets/{worksheet_id}/export-docx")
async def export_worksheet_docx(worksheet_id: str, user: User = Depends(require_user)):
    ws = await db.worksheets.find_one({"worksheet_id": worksheet_id}, {"_id": 0})
    if not ws:
        raise HTTPException(status_code=404, detail="Worksheet not found")
    if ws.get("user_id") != user.user_id and not ws.get("is_public", False):
        raise HTTPException(status_code=403, detail="This worksheet is private")

    content = ws.get("content", {})
    title = content.get("title", ws.get("title", "Untitled Worksheet"))

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    heading = document.add_heading(title, 0)
    heading.alignment = 1

    if content.get("subtitle"):
        sub = document.add_paragraph(content["subtitle"])
        sub.alignment = 1

    document.add_paragraph("Name: ____________________\t\tDate: __________\t\tScore: _____/100")

    def add_question(doc, q, num):
        q_text = q if isinstance(q, str) else q.get("question", q.get("sentence", q.get("prompt", q.get("text", str(q)))))
        doc.add_paragraph(f"{num}. {q_text}")
        options = []
        if isinstance(q, dict):
            options = q.get("options", q.get("choices", []))
        if options:
            for opt_idx, opt in enumerate(options):
                opt_text = opt if isinstance(opt, str) else opt.get("text", opt.get("label", str(opt)))
                opt_label = opt.get("label", chr(65 + opt_idx)) if isinstance(opt, dict) and opt.get("label") else chr(65 + opt_idx)
                doc.add_paragraph(f"    ( ) {opt_label}. {opt_text}")
        else:
            doc.add_paragraph("\n    __________________________________________________\n")

    passage = content.get("reading_passage") or content.get("passage")
    if passage:
        document.add_heading("Reading Passage", level=1)
        if isinstance(passage, dict):
            if passage.get("title"):
                document.add_heading(passage["title"], level=2)
            document.add_paragraph(passage.get("text", ""))
        else:
            document.add_paragraph(str(passage))

    vocab = content.get("vocabulary")
    if vocab:
        document.add_heading("Vocabulary", level=1)
        for item in vocab.get("glossary", []):
            document.add_paragraph(f"{item.get('word', '')} — {item.get('definition', '')}")
        for ex_idx, ex in enumerate(vocab.get("exercises", [])):
            instructions = ex.get("instructions", ex.get("prompt", f"Exercise {ex_idx+1}"))
            document.add_heading(instructions, level=2)
            for i, q in enumerate(ex.get("items", ex.get("questions", []))):
                add_question(document, q, i+1)

    comp = content.get("comprehension")
    if comp and comp.get("exercises"):
        document.add_heading("Comprehension", level=1)
        for ex_idx, ex in enumerate(comp.get("exercises", [])):
            instructions = ex.get("instructions", ex.get("prompt", f"Exercise {ex_idx+1}"))
            document.add_heading(instructions, level=2)
            for i, q in enumerate(ex.get("items", ex.get("questions", []))):
                add_question(document, q, i+1)

    grammar = content.get("grammar")
    if grammar:
        focus = grammar.get('focus', '')
        document.add_heading(f"Grammar{': ' + focus if focus else ''}", level=1)
        if grammar.get("explanation"):
            document.add_paragraph(grammar["explanation"])
        for ex_idx, ex in enumerate(grammar.get("exercises", [])):
            instructions = ex.get("instructions", ex.get("prompt", f"Exercise {ex_idx+1}"))
            document.add_heading(instructions, level=2)
            for i, q in enumerate(ex.get("items", ex.get("questions", []))):
                add_question(document, q, i+1)

    exercises = content.get("exercises")
    if exercises:
        for ex_idx, ex in enumerate(exercises):
            instructions = ex.get("instructions", ex.get("prompt", f"Exercise {ex_idx+1}"))
            document.add_heading(instructions, level=1)
            for i, q in enumerate(ex.get("items", ex.get("questions", []))):
                add_question(document, q, i+1)

    sections = content.get("sections")
    if sections:
        for sec_idx, sec in enumerate(sections):
            stitle = sec.get("section_title", f"Section {sec_idx+1}")
            document.add_heading(stitle, level=1)
            if sec.get("instructions"):
                document.add_paragraph(sec["instructions"])
            for q in sec.get("questions", []):
                num = q.get("number", "?")
                add_question(document, q, num)

    writing = content.get("writing") or content.get("writing_task")
    if writing:
        document.add_heading("Writing Task", level=1)
        if isinstance(writing, dict):
            task = writing.get("task") or writing.get("prompt")
            if task:
                document.add_paragraph(task)
        document.add_paragraph("\n")
        for _ in range(8):
            document.add_paragraph("_________________________________________________________________________________________")

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    headers = {'Content-Disposition': f'attachment; filename="{title}.docx"'}
    return Response(
        content=file_stream.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )

@api_router.patch("/worksheets/{worksheet_id}")
async def update_worksheet(worksheet_id: str, payload: UpdateWorksheetRequest, user: User = Depends(require_user)):
    ws = await db.worksheets.find_one({"worksheet_id": worksheet_id, "user_id": user.user_id})
    if not ws:
        raise HTTPException(status_code=404, detail="Worksheet not found")
    updates = {k: v for k, v in {"title": payload.title, "content": payload.content, "is_public": payload.is_public}.items() if v is not None}
    if updates:
        updates["updated_at"] = _now().isoformat()
        await db.worksheets.update_one({"worksheet_id": worksheet_id}, {"$set": updates})
    return {"status": "updated"}

# ============================================================
# API ROUTES — AI EDITOR
# ============================================================
@api_router.post("/worksheets/ai-edit")
async def ai_edit_worksheet(payload: AIEditRequest, user: User = Depends(require_user)):
    user = await refresh_user_credits(user)
    if user.ai_edit_credits < 1:
        raise HTTPException(status_code=402, detail="No AI edit credits remaining. You get 1 free edit per month, or unlimited with Premium.")

    ws = await db.worksheets.find_one({"worksheet_id": payload.worksheet_id, "user_id": user.user_id})
    if not ws:
        raise HTTPException(status_code=404, detail="Worksheet not found")

    edit_prompt = (
        f"You are editing an ESL worksheet.\nCurrent content:\n{json.dumps(ws['content'], indent=2)}\n\n"
        f"TEACHER'S REQUEST: {payload.command}\n\n"
        "Rules:\n- Return FULL updated worksheet as valid JSON.\n"
        "- Preserve existing structure.\n- Only modify what was requested.\n- OUTPUT MUST BE RAW JSON ONLY."
    )

    config = TIER_CONFIG.get(user.subscription_tier, TIER_CONFIG["free"])
    edited = await _run_gemini(edit_prompt, ws["level"], model_name=config["model"])
    await db.users.update_one({"user_id": user.user_id}, {"$inc": {"ai_edit_credits": -1}})

    new_id  = f"ws_{uuid.uuid4().hex[:12]}"
    new_doc = {
        "worksheet_id": new_id, "user_id": user.user_id,
        "title": f"{ws['title']} (AI Edited)",
        "level": ws["level"], "cefr": ws["cefr"], "skill": ws["skill"],
        "topic": ws.get("topic", ""), "content": edited,
        "is_public": False, "parent_id": payload.worksheet_id,
        "edit_command": payload.command, "created_at": _now().isoformat()
    }
    await db.worksheets.insert_one(new_doc)
    return {k: v for k, v in new_doc.items() if k != "_id"}

@api_router.post("/worksheets/fix")
async def fix_worksheet(payload: WorksheetFixRequest, user: User = Depends(require_user)):
    ws = await db.worksheets.find_one({"worksheet_id": payload.worksheetId, "user_id": user.user_id})
    if not ws:
        raise HTTPException(status_code=404, detail="Worksheet not found")

    fix_prompt = (
        f"You are fixing an ESL worksheet.\nCurrent content:\n{json.dumps(ws.get('content', {}), indent=2)}\n\n"
        f"Original prompt/context:\n{payload.originalPrompt or 'Not provided'}\n\n"
        f"Teacher feedback:\n{payload.feedback}\n\n"
        "Return the FULL corrected worksheet content as raw valid JSON only. Preserve the existing schema where possible."
    )
    config = TIER_CONFIG.get(user.subscription_tier, TIER_CONFIG["free"])
    fixed = await _run_gemini(fix_prompt, ws.get("level", "Primary"), model_name=config["model"])
    await db.worksheets.update_one(
        {"worksheet_id": payload.worksheetId, "user_id": user.user_id},
        {"$set": {"content": fixed, "updated_at": _now().isoformat(), "fix_feedback": payload.feedback}}
    )
    return {"status": "fixed", "content": fixed}

# ============================================================
# API ROUTES — REWARDED ADS
# ============================================================
@api_router.post("/usage/grant-rewarded")
async def grant_sponsor_reward(payload: RewardedAdRequest, user: User = Depends(require_user)):
    if payload.reward_type == "ai_edit":
        if user.subscription_tier != "premium":
            raise HTTPException(status_code=403, detail="AI edit rewards require Premium.")
        bonus = 1 if payload.tier <= 15 else 2 if payload.tier <= 30 else 3
        await db.users.update_one({"user_id": user.user_id}, {"$inc": {"ai_edit_credits": bonus}})
        return {"status": "reward_granted", "amount": bonus, "type": "ai_edit_credit"}
    bonus = 1 if payload.tier <= 15 else 2
    await db.users.update_one({"user_id": user.user_id}, {"$inc": {"bonus_credits": bonus}})
    return {"status": "reward_granted", "amount": bonus, "type": "worksheet_credit"}

# ============================================================
# API ROUTES — PUBLIC LIBRARY
# ============================================================
@api_router.get("/library/feed")
async def public_library_feed(
    response: Response,
    level: Optional[str] = None,
    skill: Optional[str] = None,
    search: Optional[str] = None
):
    response.headers["Cache-Control"] = "public, max-age=300, stale-while-revalidate=600"
    response.headers["X-Robots-Tag"] = "index, follow"
    query: dict = {"is_public": True}
    if level  and level  != "All": query["level"] = level
    if skill  and skill  != "All": query["skill"] = skill
    if search:
        query["$or"] = [
            {"title": {"$regex": search, "$options": "i"}},
            {"topic": {"$regex": search, "$options": "i"}}
        ]
    pipeline = [
        {"$match": query}, {"$sort": {"created_at": -1}}, {"$limit": 100},
        {"$lookup": {"from": "users", "localField": "user_id", "foreignField": "user_id", "as": "author"}},
        {"$project": {
            "_id": 0, "worksheet_id": 1, "title": 1, "level": 1, "cefr": 1,
            "skill": 1, "topic": 1, "created_at": 1,
            "author_name": {"$ifNull": [{"$arrayElemAt": ["$author.name", 0]}, "Anonymous Teacher"]}
        }}
    ]
    return await db.worksheets.aggregate(pipeline).to_list(100)

@api_router.get("/library/feed.xml")
async def public_library_feed_xml(
    response: Response,
    level: Optional[str] = None,
    skill: Optional[str] = None,
    search: Optional[str] = None
):
    response.headers["Cache-Control"] = "public, max-age=300, stale-while-revalidate=600"
    response.headers["X-Robots-Tag"] = "index, follow"
    response.headers["Content-Type"] = "application/atom+xml; charset=utf-8"

    query: dict = {"is_public": True}
    if level and level != "All":
        query["level"] = level
    if skill and skill != "All":
        query["skill"] = skill
    if search:
        query["$or"] = [
            {"title": {"$regex": search, "$options": "i"}},
            {"topic": {"$regex": search, "$options": "i"}}
        ]

    pipeline = [
        {"$match": query}, {"$sort": {"created_at": -1}}, {"$limit": 100},
        {"$lookup": {"from": "users", "localField": "user_id", "foreignField": "user_id", "as": "author"}},
        {"$project": {
            "_id": 0, "worksheet_id": 1, "title": 1, "level": 1, "cefr": 1,
            "skill": 1, "topic": 1, "created_at": 1,
            "author_name": {"$ifNull": [{"$arrayElemAt": ["$author.name", 0]}, "Anonymous Teacher"]}
        }}
    ]

    worksheets = await db.worksheets.aggregate(pipeline).to_list(100)
    updated_at = worksheets[0]["created_at"] if worksheets else _now().isoformat()

    entries = []
    for ws in worksheets:
        ws_title = xml_escape(ws.get("title", "Untitled Worksheet"))
        ws_topic = ws.get("topic")
        topic_text = f" on {xml_escape(ws_topic)}" if ws_topic else ""
        summary = (
            f"ESL worksheet for {xml_escape(ws['level'])} "
            f"({xml_escape(ws['cefr'])}) — {xml_escape(ws['skill'])}{topic_text}."
        )
        url = f"https://www.smartgiaoan.site/worksheet/{xml_escape(ws['worksheet_id'])}"
        entries.append("\n".join([
            "  <entry>",
            f"    <id>{url}</id>",
            f"    <title>{ws_title}</title>",
            f"    <link href=\"{url}\" />",
            f"    <updated>{xml_escape(ws['created_at'])}</updated>",
            f"    <summary>{summary}</summary>",
            f"    <author><name>{xml_escape(ws['author_name'])}</name></author>",
            "  </entry>"
        ]))

    feed = "\n".join([
        "<?xml version=\"1.0\" encoding=\"utf-8\"?>",
        "<feed xmlns=\"http://www.w3.org/2005/Atom\">",
        "  <title>SmartGiaoAn Public Library Feed</title>",
        "  <id>https://www.smartgiaoan.site/library</id>",
        "  <link href=\"https://www.smartgiaoan.site/library\" />",
        "  <link href=\"https://smartgiaoan.onrender.com/api/library/feed.xml\" rel=\"self\" />",
        f"  <updated>{xml_escape(updated_at)}</updated>",
        "  <subtitle>Public ESL worksheets generated by teachers across Vietnam.</subtitle>",
        *entries,
        "</feed>"
    ])

    return Response(content=feed, media_type="application/atom+xml")

@api_router.post("/library/upload")
async def upload_library_worksheet(payload: LibraryUploadRequest, user: User = Depends(require_user)):
    primary_skill = payload.skills[0] if payload.skills else "reading"
    ws_id = f"ws_{uuid.uuid4().hex[:12]}"
    ws_doc = {
        "worksheet_id": ws_id,
        "user_id": user.user_id,
        "title": payload.title.strip(),
        "level": payload.level,
        "cefr": "",
        "skill": primary_skill,
        "skills": payload.skills,
        "topic": payload.topic or "",
        "description": payload.description.strip(),
        "content": {
            "title": payload.title.strip(),
            "description": payload.description.strip(),
            "sections": [],
            "answer_key": [],
            "teacher_notes": "Community upload awaiting review.",
        },
        "is_public": bool(payload.is_public),
        "review_status": "pending",
        "created_at": _now().isoformat(),
    }
    await db.worksheets.insert_one(ws_doc)
    return {k: v for k, v in ws_doc.items() if k != "_id"}

@api_router.post("/library/{worksheet_id}/clone")
async def clone_worksheet(worksheet_id: str, user: User = Depends(require_user)):
    original = await db.worksheets.find_one({"worksheet_id": worksheet_id, "is_public": True})
    if not original:
        raise HTTPException(status_code=404, detail="Worksheet not found or not public")
    new_id  = f"ws_{uuid.uuid4().hex[:12]}"
    new_doc = {
        "worksheet_id": new_id, "user_id": user.user_id,
        "title": original["title"], "level": original["level"],
        "cefr": original["cefr"], "skill": original["skill"],
        "topic": original.get("topic", ""), "content": original["content"],
        "is_public": False, "created_at": _now().isoformat(), "cloned_from": worksheet_id
    }
    await db.worksheets.insert_one(new_doc)
    return {"worksheet_id": new_id, "status": "cloned"}

# ============================================================
# API ROUTES — BILLING
# ============================================================
@api_router.get("/billing/tier")
async def get_tier(user: User = Depends(require_user)):
    user   = await refresh_user_credits(user)
    config = TIER_CONFIG.get(user.subscription_tier, TIER_CONFIG["free"])
    # FIX 4: Use .get() with a safe fallback so free tier (which has lifetime_quota, not monthly_quota) doesn't crash
    monthly_quota = config.get("monthly_quota", config.get("lifetime_quota", 3))
    return {
        "tier": user.subscription_tier,
        "is_premium": user.is_premium,
        "monthly_quota": monthly_quota,
        "used_this_month": user.free_used,
        "remaining_this_month": (
            max(0, monthly_quota + user.bonus_credits - user.free_used)
            if user.subscription_tier != "free" else "unlimited"
        ),
        "ai_edit_credits": user.ai_edit_credits,
        "has_word_editor": config["has_word_editor"],
        "has_ai_editor":   config["has_ai_editor"],
        "has_ads":         config["has_ads"],
        "model":           config["model"],
        "ai_engine":       "vertex_ai" if USE_VERTEX_AI else "generative_ai",
        "ai_region":       GEMINI_REGION,
        "reset_at": user.monthly_reset_at.isoformat() if user.monthly_reset_at else None,
    }

@api_router.post("/billing/paypal-capture")
async def paypal_capture(payload: PayPalCaptureRequest, user: User = Depends(require_user)):
    now = _now()
    if payload.product_type in ("premium_monthly", "pro_monthly"):
        try:
            sub = await verify_paypal_subscription(payload.order_id)
            if sub.get("status") not in ("ACTIVE", "APPROVED"):
                raise HTTPException(status_code=400, detail=f"Subscription not active: {sub.get('status')}")
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(status_code=400, detail="PayPal verification failed")
        tier = "premium" if payload.product_type == "premium_monthly" else "pro"
        await db.users.update_one({"user_id": user.user_id}, {"$set": {
            "subscription_tier": tier, "is_premium": True, "free_used": 0,
            "ai_edit_credits": TIER_CONFIG[tier]["ai_edits_per_month"],
            "monthly_reset_at": (now + timedelta(days=30)).isoformat(),
            "paypal_subscription_id": payload.order_id
        }})
        return {"status": "success", "tier": tier}
    elif payload.product_type == "ai_edit_pack":
        try:
            order = await verify_paypal_order(payload.order_id)
            if order.get("status") not in ("COMPLETED", "APPROVED"):
                raise HTTPException(status_code=400, detail=f"Order not completed: {order.get('status')}")
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(status_code=400, detail="PayPal verification failed")
        await db.users.update_one({"user_id": user.user_id}, {"$inc": {"ai_edit_credits": 10}})
        return {"status": "success", "credits_added": 10}
    raise HTTPException(status_code=400, detail="Unknown product type")

@api_router.post("/billing/mark-premium")
async def mark_premium(user: User = Depends(require_user)):
    await db.users.update_one({"user_id": user.user_id}, {"$set": {
        "subscription_tier": "premium", "is_premium": True, "free_used": 0,
        "ai_edit_credits": TIER_CONFIG["premium"]["ai_edits_per_month"],
        "monthly_reset_at": (_now() + timedelta(days=30)).isoformat()
    }})
    return {"status": "premium_activated"}

@api_router.post("/billing/mark-pro")
async def mark_pro(user: User = Depends(require_user)):
    await db.users.update_one({"user_id": user.user_id}, {"$set": {
        "subscription_tier": "pro", "is_premium": True, "free_used": 0,
        "ai_edit_credits": TIER_CONFIG["pro"]["ai_edits_per_month"],
        "monthly_reset_at": (_now() + timedelta(days=30)).isoformat()
    }})
    return {"status": "pro_activated"}

@api_router.post("/billing/downgrade")
async def downgrade(user: User = Depends(require_user)):
    await db.users.update_one({"user_id": user.user_id}, {"$set": {
        "subscription_tier": "free", "is_premium": False,
        "paypal_subscription_id": None, "ai_edit_credits": 0
    }})
    return {"status": "downgraded_to_free"}

# ============================================================
# PAYPAL HELPERS
# ============================================================
async def _paypal_access_token() -> str:
    if not PAYPAL_CLIENT_ID or not PAYPAL_CLIENT_SECRET:
        raise HTTPException(status_code=500, detail="PayPal credentials not configured")
    async with httpx.AsyncClient(timeout=15.0) as client:
        auth_str = base64.b64encode(f"{PAYPAL_CLIENT_ID}:{PAYPAL_CLIENT_SECRET}".encode()).decode()
        r = await client.post(
            f"{PAYPAL_BASE_URL}/v1/oauth2/token",
            headers={"Authorization": f"Basic {auth_str}"},
            data={"grant_type": "client_credentials"}
        )
        r.raise_for_status()
        return r.json()["access_token"]

async def verify_paypal_order(order_id: str) -> dict:
    token = await _paypal_access_token()
    async with httpx.AsyncClient(timeout=15.0) as client:
        r = await client.get(
            f"{PAYPAL_BASE_URL}/v2/checkout/orders/{order_id}",
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
        )
        r.raise_for_status()
        return r.json()

async def verify_paypal_subscription(subscription_id: str) -> dict:
    token = await _paypal_access_token()
    async with httpx.AsyncClient(timeout=15.0) as client:
        r = await client.get(
            f"{PAYPAL_BASE_URL}/v1/billing/subscriptions/{subscription_id}",
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
        )
        r.raise_for_status()
        return r.json()

# ============================================================
# API ROUTES — PAYPAL WEBHOOKS
# ============================================================
@api_router.post("/webhooks/paypal")
async def paypal_webhook(request: Request):
    try:
        paypal_cert_url         = request.headers.get("paypal-transmission-certurl")
        paypal_transmission_id  = request.headers.get("paypal-transmission-id")
        paypal_transmission_time= request.headers.get("paypal-transmission-time")
        paypal_auth_algo        = request.headers.get("paypal-auth-algo")
        paypal_transmission_sig = request.headers.get("paypal-transmission-sig")
        paypal_webhook_id       = os.environ.get("PAYPAL_WEBHOOK_ID")

        if not all([paypal_cert_url, paypal_transmission_id, paypal_transmission_time,
                    paypal_auth_algo, paypal_webhook_id, paypal_transmission_sig]):
            logger.warning("Missing PayPal webhook headers or WEBHOOK_ID.")
            raise HTTPException(status_code=400, detail="Missing PayPal webhook headers or WEBHOOK_ID.")

        body = await request.body()

        verification_data = {
            "transmission_id":   paypal_transmission_id,
            "transmission_time": paypal_transmission_time,
            "cert_url":          paypal_cert_url,
            "auth_algo":         paypal_auth_algo,
            "transmission_sig":  paypal_transmission_sig,
            "webhook_id":        paypal_webhook_id,
            "webhook_event":     json.loads(body)
        }

        async with httpx.AsyncClient(timeout=15.0) as client:
            verify_url = f"{PAYPAL_BASE_URL}/v1/notifications/verify-webhook-signature"
            r = await client.post(verify_url, json=verification_data)
            r.raise_for_status()
            verification_status = r.json()

        if verification_status.get("verification_status") != "SUCCESS":
            logger.error(f"PayPal webhook verification failed: {verification_status}")
            raise HTTPException(status_code=403, detail="PayPal webhook verification failed.")

        data       = json.loads(body)
        event_type = data.get("event_type", "")
        resource   = data.get("resource", {})
        logger.info(f"PayPal webhook event received: {event_type}")

        custom_id = resource.get("custom_id") or resource.get("subscriber", {}).get("custom_id", "")
        if not custom_id or not custom_id.startswith("user_"):
            logger.warning(f"PayPal webhook received with missing or invalid custom_id: {custom_id}")
            return {"status": "ignored", "message": "Missing or invalid custom_id"}

        sub_id  = resource.get("id", "")
        plan_id = resource.get("plan_id", "")

        if event_type == "BILLING.SUBSCRIPTION.ACTIVATED":
            tier = "pro" if plan_id == os.environ.get("PAYPAL_PRO_PLAN_ID") else "premium"
            await db.users.update_one({"user_id": custom_id}, {"$set": {
                "subscription_tier": tier, "is_premium": True, "free_used": 0,
                "ai_edit_credits": TIER_CONFIG[tier]["ai_edits_per_month"],
                "monthly_reset_at": (_now() + timedelta(days=30)).isoformat(),
                "paypal_subscription_id": sub_id
            }})
            logger.info(f"User {custom_id} activated {tier} subscription {sub_id}")

        elif event_type in ("BILLING.SUBSCRIPTION.CANCELLED", "BILLING.SUBSCRIPTION.EXPIRED"):
            await db.users.update_one({"user_id": custom_id}, {"$set": {
                "subscription_tier": "free", "is_premium": False,
                "paypal_subscription_id": None, "ai_edit_credits": 0
            }})
            logger.info(f"User {custom_id} subscription {sub_id} cancelled/expired.")

        elif event_type == "BILLING.SUBSCRIPTION.SUSPENDED":
            await db.users.update_one({"user_id": custom_id}, {"$set": {
                "subscription_tier": "free", "is_premium": False, "ai_edit_credits": 0
            }})
            logger.warning(f"User {custom_id} subscription {sub_id} suspended.")

        elif event_type in ("BILLING.SUBSCRIPTION.RE-ACTIVATED", "BILLING.SUBSCRIPTION.REACTIVATED"):
            tier = "pro" if plan_id == os.environ.get("PAYPAL_PRO_PLAN_ID") else "premium"
            await db.users.update_one({"user_id": custom_id}, {"$set": {
                "subscription_tier": tier, "is_premium": True, "free_used": 0,
                "ai_edit_credits": TIER_CONFIG[tier]["ai_edits_per_month"],
                "monthly_reset_at": (_now() + timedelta(days=30)).isoformat(),
                "paypal_subscription_id": sub_id
            }})
            logger.info(f"User {custom_id} subscription {sub_id} re-activated to {tier}.")

        elif event_type == "BILLING.SUBSCRIPTION.UPDATED":
            tier = "pro" if plan_id == os.environ.get("PAYPAL_PRO_PLAN_ID") else "premium"
            await db.users.update_one({"user_id": custom_id}, {"$set": {
                "subscription_tier": tier,
                "ai_edit_credits": TIER_CONFIG[tier]["ai_edits_per_month"],
                "monthly_reset_at": (_now() + timedelta(days=30)).isoformat()
            }})
            logger.info(f"User {custom_id} subscription {sub_id} updated to {tier}.")

        elif event_type == "BILLING.SUBSCRIPTION.PAYMENT.FAILED":
            # PayPal will retry; we just log + flag the user. Keep tier active until SUSPENDED/CANCELLED.
            logger.warning(f"User {custom_id} subscription {sub_id} payment failed (PayPal will retry).")

        elif event_type == "BILLING.SUBSCRIPTION.CREATED":
            # Subscription created but not yet ACTIVATED — no DB change needed.
            logger.info(f"User {custom_id} subscription {sub_id} created (awaiting activation).")

        elif event_type == "PAYMENT.SALE.COMPLETED":
            if "-" in custom_id:
                user_id_part, product_type = custom_id.split("-", 1)
                if product_type == "ai_edit_pack":
                    await db.users.update_one({"user_id": user_id_part}, {"$inc": {"ai_edit_credits": 10}})
                    logger.info(f"User {user_id_part} received 10 AI edit credits from one-time payment.")
                else:
                    logger.warning(f"Unhandled one-time product type: {product_type} for user {user_id_part}")
            else:
                logger.warning(f"One-time payment without product_type in custom_id: {custom_id}")
        else:
            logger.info(f"Unhandled PayPal webhook event: {event_type}")

        return {"status": "ok"}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PayPal webhook error: {e}")
        return {"status": "ok"}

# NEW EMAIL VERIFICATION ENDPOINTS AND HELPER
class EmailVerificationRequest(BaseModel):
    token: str

async def _send_email(to_email: str, subject: str, content: str) -> None:
    provider = os.environ.get('EMAIL_SERVICE_PROVIDER', '')
    if provider and provider.lower() == 'sendgrid':
        key = os.environ.get('EMAIL_API_KEY') or os.environ.get('SENDGRID_API_KEY')
        sender = os.environ.get('EMAIL_FROM', 'noreply@smartgiaoan.com')
        if not key:
            logger.warning('SendGrid API key not configured; skipping email send')
            return
        payload = {
            "personalizations": [{ "to": [ { "email": to_email } ] }],
            "from": { "email": sender },
            "subject": subject,
            "content": [ { "type": "text/plain", "value": content } ]
        }
        headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
        async with httpx.AsyncClient() as client:
            r = await client.post("https://api.sendgrid.com/v3/mail/send", headers=headers, json=payload)
            try:
                r.raise_for_status()
            except Exception:
                logger.warning("Email send failed via SendGrid")
    else:
        logger.info("Email service not configured; would send verification to: %s", to_email)

@api_router.post("/auth/send-verification")
async def send_verification(user: User = Depends(require_user)):
    if getattr(user, 'email_verified', False):
        return {"status": "already_verified"}
    secret = os.environ.get('EMAIL_VERIFICATION_JWT_SECRET') or os.environ.get('JWT_VERIFICATION_SECRET')
    if not secret:
        raise HTTPException(status_code=500, detail="Email verification secret not configured")
    exp = datetime.utcnow() + timedelta(hours=24)
    data = {'user_id': user.user_id, 'email': user.email, 'exp': exp.isoformat()}
    token = jwt.encode(data, secret, algorithm='HS256')
    link = f"{FRONTEND_URL}/verify-email?token={token}"
    await _send_email(user.email, 'Verify your SmartGiaoAn email', f'Please verify: {link}')
    return {"status": "verification_sent"}

class EmailVerificationRequest(BaseModel):
    token: str

@api_router.post("/auth/verify-email")
async def verify_email(payload: EmailVerificationRequest):
    secret = os.environ.get('EMAIL_VERIFICATION_JWT_SECRET') or os.environ.get('JWT_VERIFICATION_SECRET')
    if not secret:
        raise HTTPException(status_code=500, detail="Email verification secret not configured")
    try:
        data = jwt.decode(payload.token, secret, algorithms=["HS256"])
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid verification token")
    user_id = data.get("user_id")
    if not user_id:
        raise HTTPException(status_code=400, detail="Invalid token payload")
    await db.users.update_one({"user_id": user_id}, {"$set": {"email_verified": True}})
    return {"status": "verified"}
# MIDDLEWARE & ROUTING
# ============================================================
app.include_router(api_router)
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

@app.get("/")
async def root():
    return {
        "app": "SmartGiaoAn API", "status": "operational", "version": "3.3.0",
        "ai_engine": "vertex_ai" if USE_VERTEX_AI else "generative_ai",
        "ai_region": GEMINI_REGION,
    }

@app.head("/")
async def root_head():
    return Response(status_code=200)

@app.get("/worksheet_seo/{worksheet_id}")
async def worksheet_seo_data(worksheet_id: str):
    ws = await db.worksheets.find_one({"worksheet_id": worksheet_id}, {"_id": 0})
    if not ws:
        raise HTTPException(status_code=404, detail="Worksheet not found")
    return ws

@app.get("/health")
async def health():
    return {
        "status": "healthy",
        "db": "connected" if mongo_client else "disconnected",
        "ai_engine": "vertex_ai" if USE_VERTEX_AI else "generative_ai",
        "ai_region": GEMINI_REGION,
    }